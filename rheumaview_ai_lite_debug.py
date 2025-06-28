import streamlit as st
import torch
from torchvision import transforms
from PIL import Image
from inference_core import predict_region
from docx import Document
from datetime import datetime
from io import BytesIO

CLASS_NAMES = [
    "Cervical Spine", "Thoracic Spine", "Lumbar Spine", "Pelvis/SI Joints/ Sacrum", 
    "Hips", "Knees", "Ankles/Feet", 
    "Shoulders", "Elbows", "Hands/Wrists", "Long bones"
]

st.set_page_config(page_title="RheumaView Lite", layout="centered")
st.title("🟩 RheumaView™ Lite")
st.caption("Curated by Dr. Olga Goodman • Region classifier demo")

uploaded_files = st.file_uploader("Upload X-ray images", accept_multiple_files=True, type=["jpg", "jpeg", "png", "bmp", "tiff", "webp"])

results = []

if uploaded_files:
    st.markdown(f"**Total files uploaded:** {len(uploaded_files)}")
    st.subheader("Image Preview, Prediction, and Manual Override")

    enable_ai_prediction = st.checkbox("Enable AI prediction (uncheck to use manual only)", value=True)

    for file in uploaded_files:
        image = Image.open(file).convert("L")
        image_rgb = image.convert("RGB")
        col1, col2 = st.columns([1, 2])
        with col1:
            st.image(image, caption=f"{file.name}", width=120)
        with col2:
            if enable_ai_prediction:
                try:
                    top3 = predict_region(image_rgb)
                    if isinstance(top3, list) and top3 and isinstance(top3[0], (list, tuple)) and isinstance(top3[0][0], int):
                        top_label = CLASS_NAMES[top3[0][0]]
                        st.markdown(f"**Top prediction:** {top_label}")
                        st.markdown("**Confidence breakdown:**")
                        for idx, prob in top3:
                            st.markdown(f"- {CLASS_NAMES[idx]}: {prob:.2%}")
                        manual = st.selectbox(
                            f"Override region for {file.name}",
                            CLASS_NAMES,
                            index=top3[0][0],
                            key=f"{file.name}_pred"
                        )
                    else:
                        raise ValueError("Unexpected prediction output format")
                except Exception as e:
                    st.error(f"Prediction failed for {file.name}: {e}")
                    manual = st.selectbox(
                        f"Manual region for {file.name} (prediction failed)",
                        CLASS_NAMES,
                        key=f"{file.name}_failed"
                    )
            else:
                manual = st.selectbox(
                    f"Manual region for {file.name}",
                    CLASS_NAMES,
                    key=f"{file.name}_manual"
                )

        results.append((file.name, manual))

    # EMR Summary
    if st.button("Generate EMR Summary"):
        emr_text = []
        region_count = {}
        for _, region in results:
            region_count[region] = region_count.get(region, 0) + 1

        for region, count in region_count.items():
            emr_text.append(f"{region} – {count} view(s)")

        emr_summary = "Study includes: " + "; ".join(emr_text) + "."
        st.success("EMR Summary:")
        st.code(emr_summary, language="markdown")

    # DOCX Report
    if st.button("Generate Report (.docx)"):
        doc = Document()
        doc.add_heading("RheumaView™ Radiology Report", level=1)
        doc.add_paragraph("Curated by Dr. Olga Goodman")
        doc.add_paragraph(f"Date: {datetime.today().strftime('%Y-%m-%d')}")

        region_groups = {}
        for file_name, region in results:
            region_groups.setdefault(region, []).append(file_name)

        for region, files in region_groups.items():
            doc.add_heading(region, level=2)
            for fname in files:
                doc.add_paragraph(f"- {fname}")

        output_stream = BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)

        st.success("Report generated successfully!")
        st.download_button(
            label="Download Report",
            file_name="rheumaview_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            data=output_stream.read()
        )
